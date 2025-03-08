import os
import glob
import sys
import re
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from pdf2docx import Converter

DISCIPLINE_MAP = {
    'Администрирование сетей передачи информации': 'АСПИ',
    'Администрирование операционных систем': 'АОС',
    'Безопасность операционных систем': 'БОС',
    'Современные операционные системы': 'СОС',
    'Разработка корпоративных дистрибутивов': 'РКД',
    'Специализированные языки и технологии программирования': 'СЯиТП'
}

def convert_pdf_to_docx(directory):
    """Offer PDF to DOCX conversion with original filename"""
    # Find all PDF files in directory
    pdf_files = glob.glob(os.path.join(directory, '*.pdf'))
    
    if not pdf_files:
        print("\nВ директории не найдено PDF-файлов.")
        return
    
    # Display available PDF files
    print("\nНайдены PDF-файлы:")
    for idx, pdf in enumerate(pdf_files, 1):
        print(f"{idx}. {os.path.basename(pdf)}")
    print("0. Не конвертировать файлы")
    
    while True:
        try:
            choice = input("\nВыберите PDF для конвертации (0 чтобы пропустить): ").strip()
            if choice == '0':
                print("Конвертация отменена.")
                return
            
            choice = int(choice)
            if 1 <= choice <= len(pdf_files):
                selected_pdf = pdf_files[choice-1]
                docx_path = os.path.splitext(selected_pdf)[0] + '.docx'
                
                try:
                    cv = Converter(selected_pdf)
                    cv.convert(docx_path)
                    cv.close()
                    print(f"✅ Успешно конвертирован: {os.path.basename(docx_path)}")
                    return docx_path
                except Exception as e:
                    print(f"❌ Ошибка конвертации: {str(e)}")
                    return None
            else:
                print(f"Пожалуйста, введите число от 0 до {len(pdf_files)}")
        except ValueError:
            print("Некорректный ввод! Введите число.")
            
def process_images(directory):
    image_files = sorted(
        [f for f in glob.glob(os.path.join(directory, "*.*")) 
         if f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif'))],
        key=lambda x: os.path.basename(x).lower()
    )
    
    if not image_files:
        print("\n❌ Ошибка: изображений нет в директории!")
        print("Поддерживаемые форматы: PNG, JPG/JPEG, GIF")
        return None
    
    figure_data = []
    skipped_counter = 0
    
    for idx, img_path in enumerate(image_files, 1):
        try:
            with Image.open(img_path) as img:
                # Show image using default viewer
                img.show()
                user_input = input(f"\nИзображение {idx}/{len(image_files)}\n"
                                 f"Файл: {os.path.basename(img_path)}\n"
                                 "Введите название (описание) или '-' для пропуска: ").strip()
                
                if user_input == "-":
                    print(f"Изображение {os.path.basename(img_path)} пропущено.")
                    skipped_counter += 1
                    continue
                
                adjusted_idx = idx - skipped_counter
                
                if user_input:
                    stripped = user_input.rstrip('.')
                    if stripped:
                        processed = stripped[0].upper() + stripped[1:]
                    else:
                        processed = f"Рис. {adjusted_idx}. "
                    title = f"{processed}."
                else:
                    title = f"Рис. {adjusted_idx}. "
                
                figure_data.append({
                    'path': img_path,
                    'title': title
                })
        except Exception as e:
            print(f"\n⚠️ Ошибка обработки {os.path.basename(img_path)}: {str(e)}")
            continue
    
    return figure_data

def find_word_template(directory):
    templates = glob.glob(os.path.join(directory, "*.doc*"))
    
    if not templates:
        print("\n❌ Ошибка: в директории нет doc(x) файлов!")
        print("Добавьте файл-шаблон - template.docx")
        return None
    
    if len(templates) == 1:
        return templates[0]
    
    print("\nВыберите файл-шаблон:")
    for i, path in enumerate(templates, 1):
        print(f"{i}. {os.path.basename(path)}")
    
    while True:
        try:
            choice = int(input("\nВведите номер файла: "))
            if 1 <= choice <= len(templates):
                return templates[choice-1]
            print("Номер неверный, введите снова.")
        except ValueError:
            print("Введите число!")
            
def apply_document_styles(doc):
    """Apply all document styles including page formatting"""
    styles = doc.styles
    
    # ===== Normal text style =====
    normal_style = styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(14)
    normal_style.font.color.rgb = None  # Automatic (black)
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal_style.paragraph_format.space_after = Pt(0)  # No extra space after paragraphs
        
def get_output_filename(metadata):
    """Generate filename using metadata"""
    # Extract required fields from metadata
    surname = metadata['last_name']
    first_initial = metadata['first_name'][0].upper() if metadata['first_name'] else ''
    patron_initial = metadata['patron_name'][0].upper() if metadata['patron_name'] else ''
    report_num = metadata['num']
    group = metadata['group']
    discipline = metadata['discipline']
    
    # Format initials
    initials = f"{first_initial}{patron_initial}" if patron_initial else first_initial
    
    # Abbreviate the discipline using DISCIPLINE_MAP
    discipline_abbr = DISCIPLINE_MAP.get(discipline, discipline)
    
    # Construct the filename
    filename = f"{surname}_{initials}_ЛР{report_num}_{group}_{discipline_abbr}.docx"
    
    invalid_chars = r'\/:*?"<>|'
    for char in invalid_chars:
        filename = filename.replace(char, '')
    
    filename = filename[:50]
    
    return filename

def get_report_metadata():
    print("\n" + "="*40)
    print(" Метаданные для отчета ".center(40, "="))
    print("="*40)
    
    name_pattern = re.compile(r'^[А-Яа-яЁё-]+$')     # Russian letters and hyphens
    group_pattern = re.compile(r'^\w{2}-\d{2}-\d{2}$')  # КХ-22-01 format
    number_pattern = re.compile(r'^\d+$')            # Only digits
    
    # Personal information
    while True:
        last_name = input("\nФамилия: ").strip()
        if last_name and name_pattern.match(last_name):
            last_name = last_name[0].upper() + last_name[1:].lower()
            break
        print("❌ Фамилия не может быть пустой и должна содержать только русские буквы!")
    
    while True:
        first_name = input("Имя: ").strip()
        if first_name and name_pattern.match(first_name):
            first_name = first_name[0].upper() + first_name[1:].lower()
            break
        print("❌ Имя не может быть пустым и должно содержать только русские буквы!")
    
    while True:
        patron_name = input("Отчество: ").strip()
        if not patron_name or name_pattern.match(patron_name):
            if patron_name:
                patron_name = patron_name[0].upper() + patron_name[1:].lower()
            break
        print("❌ Отчество должно содержать только русские буквы (оставьте пустым, если нет)!")
    
    while True:
        group = input("Группа (в формате Кх-2х-хх): ").strip().upper()
        if group and group_pattern.match(group):
            break
        print("❌ Группа должна быть в формате КХ-22-01!")
    
    # Report information
    while True:
        report_num = input("\nНомер лабораторной работы: ").strip()
        if report_num and number_pattern.match(report_num):
            break
        print("❌ Номер должен быть числом!")
    
    while True:
        report_name = input("Название лабораторной работы: ").strip()
        if report_name:
            report_name = report_name[0].upper() + report_name[1:].lower()
            break
        print("❌ Название не может быть пустым!")
    
    # Discipline selection
    while True:
        print("\nВыберите дисциплину:")
        print("1. Администрирование сетей передачи информации")
        print("2. Администрирование операционных систем")
        print("3. Безопасность операционных систем")
        print("4. Современные операционные системы")
        print("5. Разработка корпоративных дистрибутивов")
        print("6. Специализированные языки и технологии программирования")
        
        discipline_choice = input("Ваш выбор (1-6): ").strip()
        
        if discipline_choice in ("1", "2", "3", "4", "5", "6"):
            break
        print("❌ Некорректный выбор! Введите число от 1 до 6.")
    
    # Map discipline choice
    if discipline_choice == "1":
        discipline = "Администрирование сетей передачи информации"
    elif discipline_choice == "2":
        discipline = "Администрирование операционных систем"
    elif discipline_choice == "3":
        discipline = "Безопасность операционных систем"
    elif discipline_choice == "4":
        discipline = "Современные операционные системы"
    elif discipline_choice == "5":
        discipline = "Разработка корпоративных дистрибутивов"
    elif discipline_choice == "6":
        discipline = "Специализированные языки и технологии программирования"
    
    return {
        "last_name": last_name,
        "first_name": first_name,
        "patron_name": patron_name,
        "group": group,
        "num": report_num,
        "name": report_name,
        "discipline": discipline
    }
    
def replace_placeholders(doc, replacements):
    """Replace placeholders in Word document"""
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
    
    # Handle tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, value)

def generate_report(directory, template_path, figure_data, output_filename, metadata):
    """Generate formatted report document with dynamic placeholders"""
    try:
        doc = Document(template_path)
        
        # Replace placeholders
        replacements = {
            "%DISCIPLINE%": metadata["discipline"],
            "%NUM%": metadata["num"],
            "%REPORT_NAME%": metadata["name"],
            "%LAST_NAME%": metadata["last_name"],
            "%FIRST_NAME%": metadata["first_name"],
            "%PATRON_NAME%": metadata.get("patron_name", ""),  # Optional field
            "%GROUP%": metadata["group"]
        }
        replace_placeholders(doc, replacements)
        
        apply_document_styles(doc)

        doc.add_section()

        execution_header = doc.add_heading('Выполнение лабораторной работы', level=1)
        execution_header.runs[0].bold = True
        doc.add_paragraph() 

        for i, item in enumerate(figure_data, 1):
            para = doc.add_paragraph(style='Normal')
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            try:
                run = para.add_run()
                run.add_picture(item['path'], width=Inches(6))
            except Exception as e:
                print(f"⚠️ Ошибка с вставкой изображения: {os.path.basename(item['path'])}")
                continue
            
            # Add caption
            caption = doc.add_paragraph(
                f"Рис. {i}. {item['title']}", 
                style='Normal'
            )
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()
        
        answers_header = doc.add_heading('Ответы на вопросы', level=1)
        answers_header.runs[0].bold = True
        conclusion_header = doc.add_heading('Заключение', level=1)
        conclusion_header.runs[0].bold = True

        output_path = os.path.join(directory, output_filename)
        doc.save(output_path)
        return output_path
        
    except Exception as e:
        print(f"\n❌ Критическая ошибка с созданием отчета: {str(e)}")
        print(f"\nПроверьте не открыт ли файл с таким же названием")
        return None
    
def main():
    report_metadata = get_report_metadata()
    print("\n" + "="*40)
    print(" Создание отчетов".ljust(39) + "=")
    print("="*40 + "\n")
    
    while True:
        directory = input("Введите абсолютный путь до изображений с шаблоном: ").strip()
        if os.path.isdir(directory):
            break
        print("Неверная директория! Введите снова.")
    
    template_path = find_word_template(directory)
    if not template_path:
        sys.exit(1)
    
    figure_data = process_images(directory)
    if not figure_data:
        sys.exit(1)
    
    output_filename = get_output_filename(report_metadata)
    
    report_path = generate_report(directory, template_path, figure_data, output_filename, report_metadata)    
    if report_path:
        print("\n" + "="*40)
        print(f"✅ Отчет сгенерировани:\n{report_path}")
        print("="*40 + "\n")
    else:
        print("\n❌ Ошибка с генерацией отчета")
        sys.exit(1)
    convert_pdf_to_docx(directory)

if __name__ == "__main__":
    main()
